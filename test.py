"""
XPIN Article DOCX to MD Converter
测试模块
"""

import unittest
import tempfile
import os
from pathlib import Path
from docx import Document
from converter import DocxToMdConverter
from config import Config
from utils import TextProcessor

class TestDocxToMdConverter(unittest.TestCase):
    """转换器测试类"""
    
    def setUp(self):
        """测试初始化"""
        self.converter = DocxToMdConverter()
        self.temp_dir = tempfile.mkdtemp()
        
    def tearDown(self):
        """测试清理"""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def create_test_docx(self, content, filename="test.docx"):
        """创建测试用的DOCX文件"""
        doc = Document()
        
        # 添加标题
        doc.add_heading('测试文档', 0)
        
        # 添加段落
        for paragraph_text in content:
            if paragraph_text.startswith('#'):
                # 添加标题
                level = paragraph_text.count('#')
                title_text = paragraph_text.lstrip('#').strip()
                doc.add_heading(title_text, min(level, 6))
            elif paragraph_text.startswith('**') and paragraph_text.endswith('**'):
                # 添加加粗段落
                p = doc.add_paragraph()
                run = p.add_run(paragraph_text[2:-2])
                run.bold = True
            else:
                # 添加普通段落
                doc.add_paragraph(paragraph_text)
        
        # 保存文件
        file_path = os.path.join(self.temp_dir, filename)
        doc.save(file_path)
        return file_path
    
    def test_basic_conversion(self):
        """测试基本转换功能"""
        test_content = [
            "这是一个测试段落",
            "这是另一个段落",
            "这是第三个段落"
        ]
        
        docx_file = self.create_test_docx(test_content)
        result = self.converter.convert_file(docx_file)
        
        self.assertTrue(result['success'])
        self.assertTrue(os.path.exists(result['output_file']))
        
        # 检查输出内容
        with open(result['output_file'], 'r', encoding='utf-8') as f:
            output_content = f.read()
        
        self.assertIn("测试文档", output_content)
        self.assertIn("这是一个测试段落", output_content)
    
    def test_heading_conversion(self):
        """测试标题转换"""
        test_content = [
            "# 一级标题",
            "## 二级标题",
            "### 三级标题",
            "正文内容"
        ]
        
        docx_file = self.create_test_docx(test_content)
        result = self.converter.convert_file(docx_file)
        
        self.assertTrue(result['success'])
        
        with open(result['output_file'], 'r', encoding='utf-8') as f:
            output_content = f.read()
        
        self.assertIn("# 一级标题", output_content)
        self.assertIn("## 二级标题", output_content)
        self.assertIn("### 三级标题", output_content)
    
    def test_bold_text_conversion(self):
        """测试加粗文本转换"""
        test_content = [
            "**这是加粗文本**",
            "普通文本",
            "**另一个加粗**"
        ]
        
        docx_file = self.create_test_docx(test_content)
        result = self.converter.convert_file(docx_file)
        
        self.assertTrue(result['success'])
        
        with open(result['output_file'], 'r', encoding='utf-8') as f:
            output_content = f.read()
        
        self.assertIn("**这是加粗文本**", output_content)
        self.assertIn("**另一个加粗**", output_content)
    
    def test_remove_patterns(self):
        """测试删除指定模式"""
        test_content = [
            "## 原文",
            "",
            "正文内容",
            "123",
            "",
            "更多内容",
            "123123",
            "结尾内容"
        ]
        
        docx_file = self.create_test_docx(test_content)
        result = self.converter.convert_file(docx_file)
        
        self.assertTrue(result['success'])
        
        with open(result['output_file'], 'r', encoding='utf-8') as f:
            output_content = f.read()
        
        # 检查"原文"和"123"是否被删除
        self.assertNotIn("## 原文", output_content)
        self.assertNotIn("123", output_content.split("正文内容")[1])  # 确保正文后的123被删除
        self.assertNotIn("123123", output_content)
        
        # 检查正常内容是否保留
        self.assertIn("正文内容", output_content)
        self.assertIn("更多内容", output_content)
        self.assertIn("结尾内容", output_content)
    
    def test_batch_conversion(self):
        """测试批量转换"""
        # 创建多个测试文件
        test_files = []
        for i in range(3):
            content = [f"文件{i+1}的内容"]
            filename = f"test{i+1}.docx"
            docx_file = self.create_test_docx(content, filename)
            test_files.append(docx_file)
        
        results = self.converter.batch_convert(self.temp_dir)
        
        self.assertEqual(len(results), 3)
        self.assertTrue(all(r['success'] for r in results))
    
    def test_unsupported_file(self):
        """测试不支持的文件格式"""
        txt_file = os.path.join(self.temp_dir, "test.txt")
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write("测试内容")
        
        result = self.converter.convert_file(txt_file)
        self.assertFalse(result['success'])
        self.assertIn('error', result)
    
    def test_nonexistent_file(self):
        """测试不存在的文件"""
        nonexistent_file = os.path.join(self.temp_dir, "nonexistent.docx")
        result = self.converter.convert_file(nonexistent_file)
        self.assertFalse(result['success'])
        self.assertIn('error', result)

class TestTextProcessor(unittest.TestCase):
    """文本处理器测试类"""
    
    def setUp(self):
        self.processor = TextProcessor()
    
    def test_remove_unwanted_patterns(self):
        """测试删除不需要的模式"""
        text = """## 原文

这是正文内容
123

更多内容
123123
结尾"""
        
        result = self.processor.remove_unwanted_patterns(text, Config.REMOVE_PATTERNS)
        
        # 检查指定模式是否被删除
        self.assertNotIn("## 原文", result)
        self.assertNotIn("123\n\n", result)
        self.assertNotIn("123123", result)
        
        # 检查正常内容是否保留
        self.assertIn("这是正文内容", result)
        self.assertIn("更多内容", result)
        self.assertIn("结尾", result)
    
    def test_clean_text(self):
        """测试文本清理"""
        text = "第一行\n\n\n第二行\n\n\n\n第三行"
        
        result = self.processor.clean_text(text)
        
        # 检查多余的空行是否被清理
        lines = result.split('\n')
        self.assertEqual(lines.count(''), 2)  # 应该有2个空行分隔3行内容

def run_tests():
    """运行所有测试"""
    print("开始运行单元测试...")
    
    # 创建测试套件
    test_suite = unittest.TestSuite()
    
    # 添加测试类
    test_suite.addTest(unittest.makeSuite(TestDocxToMdConverter))
    test_suite.addTest(unittest.makeSuite(TestTextProcessor))
    
    # 运行测试
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(test_suite)
    
    return result.wasSuccessful()

def create_sample_files():
    """创建示例测试文件"""
    print("\n创建示例测试文件...")
    
    # 创建测试目录
    test_dir = "sample_test"
    os.makedirs(test_dir, exist_ok=True)
    
    # 创建示例DOCX文件
    doc = Document()
    doc.add_heading('原文', 0)
    doc.add_heading('测试文档', 1)
    
    p1 = doc.add_paragraph()
    p1.add_run('这是').bold = False
    p1.add_run('加粗文本').bold = True
    p1.add_run('测试').bold = False
    
    doc.add_paragraph('这是普通段落')
    
    p2 = doc.add_paragraph()
    run = p2.add_run('另一段加粗文本')
    run.bold = True
    
    doc.add_paragraph('123')
    doc.add_paragraph('更多内容')
    doc.add_paragraph('123123')
    doc.add_paragraph('结尾文本')
    
    sample_file = os.path.join(test_dir, "sample.docx")
    doc.save(sample_file)
    
    print(f"示例文件已创建: {sample_file}")
    print(f"运行命令测试: python main.py {test_dir}")

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='测试工具')
    parser.add_argument('--test', action='store_true', help='运行单元测试')
    parser.add_argument('--sample', action='store_true', help='创建示例测试文件')
    parser.add_argument('--all', action='store_true', help='运行所有操作')
    
    args = parser.parse_args()
    
    if args.all or args.test:
        success = run_tests()
        if not success:
            exit(1)
    
    if args.all or args.sample:
        create_sample_files()
    
    if not any([args.test, args.sample, args.all]):
        # 默认运行单元测试
        run_tests()
